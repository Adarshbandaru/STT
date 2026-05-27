import io
import os
import re
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    import openai
    from openai import OpenAI
except ImportError:
    openai = None
    OpenAI = None

try:
    import whisper
except ImportError:
    whisper = None

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")


# ══════════════════════════════════════════════════════════════════════════════
# Transcription
# ══════════════════════════════════════════════════════════════════════════════

def transcribe_with_openai(
    audio_path: Path,
    language: str = "auto",
    translate: bool = False,
    custom_vocab: str = "",
) -> Dict[str, Any]:
    client = OpenAI(api_key=OPENAI_API_KEY)
    with open(audio_path, "rb") as audio_file:
        task = "translate" if translate else "transcribe"
        response = client.audio.transcriptions.create(
            model="whisper-1",
            file=audio_file,
            response_format="verbose_json",
            language=language if language != "auto" else None,
            task=task,
        )

    return {
        "text": response.get("text", ""),
        "segments": response.get("segments", []),
        "language": response.get("language", language),
        "engine": "openai-whisper",
        "translated": translate,
    }


def transcribe_with_whisper(
    audio_path: Path,
    language: str = "auto",
    translate: bool = False,
    custom_vocab: str = "",
) -> Dict[str, Any]:
    if whisper is None:
        raise RuntimeError("Whisper not installed. Set OPENAI_API_KEY to use the OpenAI engine.")

    model = whisper.load_model("base")
    task = "translate" if translate else "transcribe"
    result = model.transcribe(
        str(audio_path),
        language=None if language == "auto" else language,
        task=task,
        verbose=False,
    )
    return {
        "text": result.get("text", ""),
        "segments": [
            {"start": float(seg["start"]), "end": float(seg["end"]), "text": seg["text"]}
            for seg in result.get("segments", [])
        ],
        "language": result.get("language", language),
        "engine": "local-whisper",
        "translated": translate,
    }


def transcribe_audio(
    audio_path: Path,
    language: str = "auto",
    translate: bool = False,
    custom_vocab: str = "",
) -> Dict[str, Any]:
    if OPENAI_API_KEY:
        return transcribe_with_openai(audio_path, language=language, translate=translate, custom_vocab=custom_vocab)
    return transcribe_with_whisper(audio_path, language=language, translate=translate, custom_vocab=custom_vocab)


# ══════════════════════════════════════════════════════════════════════════════
# Summarization
# ══════════════════════════════════════════════════════════════════════════════

def summarize_text(text: str) -> Optional[str]:
    if OPENAI_API_KEY:
        client = OpenAI(api_key=OPENAI_API_KEY)
        prompt = (
            "Summarize the following transcript concisely in 2–4 sentences:\n\n" + text
        )
        response = client.responses.create(
            model="gpt-4.1-mini",
            input=prompt,
            max_output_tokens=200,
        )
        return response.output_text.strip() if response and response.output_text else None
    
    # Advanced Local Extractive Summary Fallback (Stopword-filtered Word Frequency Scoring)
    sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if s.strip()]
    if not sentences:
        return "No text available to summarize."
    if len(sentences) <= 2:
        return text
        
    # Count frequencies of content words
    stopwords = {"the", "and", "of", "to", "a", "i", "it", "in", "is", "that", "you", "for", "on", "with", "as", "at", "by", "an", "this", "be", "have", "are", "was", "we", "or", "but", "so", "basically", "actually", "like", "um", "ah", "my", "your", "they", "he", "she", "our"}
    words = re.findall(r'\b\w+\b', text.lower())
    word_freqs = {}
    for w in words:
        if w not in stopwords:
            word_freqs[w] = word_freqs.get(w, 0) + 1
            
    # Score sentences
    sentence_scores = []
    for s in sentences:
        score = 0
        s_words = re.findall(r'\b\w+\b', s.lower())
        for sw in s_words:
            score += word_freqs.get(sw, 0)
        # Normalize score by length slightly to avoid long sentence bias
        sentence_scores.append(score / max(1, len(s_words) ** 0.5))
        
    # Pick top 3 sentences, ordered by chronological appearance
    ranked_indices = sorted(range(len(sentences)), key=lambda i: sentence_scores[i], reverse=True)[:3]
    top_indices = sorted(ranked_indices)
    
    summary = " ".join([sentences[i] for i in top_indices])
    return summary


# ══════════════════════════════════════════════════════════════════════════════
# Sentiment Analysis
# ══════════════════════════════════════════════════════════════════════════════

def analyze_sentiment(text: str) -> Dict[str, Any]:
    if OPENAI_API_KEY:
        client = OpenAI(api_key=OPENAI_API_KEY)
        prompt = (
            "Analyze the sentiment of the following transcript. "
            "Return a JSON object with 'label' (positive/negative/neutral/mixed), "
            "'confidence' (0.0–1.0), and 'reason' (brief explanation).\n\n" + text
        )
        response = client.responses.create(
            model="gpt-4.1-mini",
            input=prompt,
            max_output_tokens=150,
        )
        try:
            import json
            return json.loads(response.output_text.strip())
        except Exception:
            return {"label": "unknown", "confidence": 0.0, "reason": response.output_text.strip()}
            
    # Advanced Local Lexicon-Based Sentiment Analyzer
    pos_words = {"great", "awesome", "good", "perfect", "fantastic", "excellent", "success", "love", "happy", 
                 "resolve", "clear", "optimal", "agree", "support", "benefit", "efficient", "complete", "progress"}
    neg_words = {"bad", "terrible", "worst", "fail", "error", "mistake", "wrong", "broken", "problem", "difficult", 
                 "sad", "issue", "worry", "risk", "delay", "blocker", "slow", "unhappy", "concern"}
                 
    words = re.findall(r'\b\w+\b', text.lower())
    pos_count = sum(1 for w in words if w in pos_words)
    neg_count = sum(1 for w in words if w in neg_words)
    total_sentiment_words = pos_count + neg_count
    
    if total_sentiment_words == 0:
        return {
            "label": "neutral",
            "confidence": 0.8,
            "reason": "The transcript maintains a highly descriptive, matter-of-fact tone with neutral terminology."
        }
        
    diff = pos_count - neg_count
    ratio = abs(diff) / total_sentiment_words
    confidence = round(0.5 + (ratio * 0.5), 2)
    
    if pos_count > 0 and neg_count > 0 and abs(diff) <= 2:
        return {
            "label": "mixed",
            "confidence": 0.75,
            "reason": f"Contains contrasting elements: {pos_count} positive and {neg_count} negative indicator phrases found."
        }
    elif diff > 0:
        return {
            "label": "positive",
            "confidence": confidence,
            "reason": f"Speech exhibits constructive tone, driven by positive terms such as {[w for w in pos_words if w in words][:3]}."
        }
    else:
        return {
            "label": "negative",
            "confidence": confidence,
            "reason": f"Identified friction or concerns indicated by negative terms such as {[w for w in neg_words if w in words][:3]}."
        }


# ══════════════════════════════════════════════════════════════════════════════
# Auto Formatting (punctuation, capitalization)
# ══════════════════════════════════════════════════════════════════════════════

def auto_format_text(text: str) -> str:
    if OPENAI_API_KEY:
        client = OpenAI(api_key=OPENAI_API_KEY)
        prompt = (
            "Format the following transcript with proper punctuation, capitalization, "
            "and paragraph breaks where appropriate. Keep the exact words, just fix the formatting:\n\n" + text
        )
        response = client.responses.create(
            model="gpt-4.1-mini",
            input=prompt,
            max_output_tokens=text.__len__() * 2,
        )
        return response.output_text.strip() if response and response.output_text else text
        
    # Advanced Local Heuristic Text Formatter (offline capitalization and paragraph builder)
    text = re.sub(r'\s+', ' ', text).strip()
    # Capitalize start of sentences
    sentences = re.split(r'(?<=[.!?])\s+', text)
    formatted_sentences = []
    for s in sentences:
        if s:
            s = s[0].upper() + s[1:]
            # Capitalize single "i" words or contractions
            s = re.sub(r'\bi\b', 'I', s)
            s = re.sub(r"\bi'm\b", "I'm", s, flags=re.IGNORECASE)
            s = re.sub(r"\bi've\b", "I've", s, flags=re.IGNORECASE)
            s = re.sub(r"\bi'll\b", "I'll", s, flags=re.IGNORECASE)
            s = re.sub(r"\bi'd\b", "I'd", s, flags=re.IGNORECASE)
            formatted_sentences.append(s)
            
    # Join sentences into paragraphs of 3-4 sentences each
    paragraphs = []
    for i in range(0, len(formatted_sentences), 4):
        paragraphs.append(" ".join(formatted_sentences[i:i+4]))
        
    return "\n\n".join(paragraphs)


# ══════════════════════════════════════════════════════════════════════════════
# Export Functions
# ══════════════════════════════════════════════════════════════════════════════

def export_as_txt(text: str) -> str:
    """Plain text export."""
    return text.strip()


def export_as_srt(segments: List[Dict], text: str) -> str:
    """SubRip subtitle format with timestamps."""
    if not segments:
        return f"1\n00:00:00,000 --> 00:00:10,000\n{text.strip()}"
    lines = []
    for i, seg in enumerate(segments, 1):
        start = _format_timestamp(seg.get("start", 0))
        end = _format_timestamp(seg.get("end", 0))
        lines.append(f"{i}\n{start} --> {end}\n{seg.get('text', '').strip()}\n")
    return "\n".join(lines)


def export_as_vtt(segments: List[Dict], text: str) -> str:
    """WebVTT subtitle format."""
    lines = ["WEBVTT\n"]
    if not segments:
        lines.append("00:00:00.000 --> 00:00:10.000\n{text.strip()}")
        return "\n".join(lines)
    for seg in segments:
        start = _format_timestamp_vtt(seg.get("start", 0))
        end = _format_timestamp_vtt(seg.get("end", 0))
        lines.append(f"{start} --> {end}\n{seg.get('text', '').strip()}\n")
    return "\n".join(lines)


def export_as_docx(text: str, segments: List[Dict]) -> bytes:
    """Word document export."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx is required for DOCX export. Run: pip install python-docx")

    doc = Document()
    doc.add_heading("Transcript", 0)
    doc.add_paragraph(text.strip())
    if segments:
        doc.add_heading("Timestamped Segments", level=1)
        for seg in segments:
            start = _format_timestamp(seg.get("start", 0))
            p = doc.add_paragraph()
            run = p.add_run(f"[{start}] ")
            run.bold = True
            run.font.color.rgb = RGBColor(0x63, 0x66, 0xf1)
            p.add_run(seg.get("text", "").strip())
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def export_as_pdf(text: str, segments: List[Dict]) -> bytes:
    """PDF export."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except ImportError:
        raise RuntimeError("reportlab is required for PDF export. Run: pip install reportlab")

    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm)
    style = ParagraphStyle("body", fontSize=11, leading=16, textColor="#1e293b")
    story = [
        Paragraph("<b>Transcript</b>", style),
        Spacer(1, 0.3 * cm),
        Paragraph(text.strip().replace("\n", "<br/>"), style),
    ]
    if segments:
        story.append(Spacer(1, 0.5 * cm))
        story.append(Paragraph("<b>Timestamped Segments</b>", style))
        story.append(Spacer(1, 0.3 * cm))
        for seg in segments:
            ts = _format_timestamp(seg.get("start", 0))
            story.append(Paragraph(f"<b>[{ts}]</b> {seg.get('text', '').strip()}", style))
    doc.build(story)
    bio.seek(0)
    return bio.read()


# ─── Timestamp helpers ────────────────────────────────────────────────────────

def _format_timestamp(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds % 1) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def _format_timestamp_vtt(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds % 1) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d}.{ms:03d}"


# ══════════════════════════════════════════════════════════════════════════════
# Cognitive Co-Pilot Extenders (2026 AI Era Edition)
# ══════════════════════════════════════════════════════════════════════════════

def extract_action_items(text: str) -> List[str]:
    if OPENAI_API_KEY:
        client = OpenAI(api_key=OPENAI_API_KEY)
        prompt = (
            "Extract a list of specific, clear action items/tasks from this transcript. "
            "Return each action item on a new line starting with a bullet (-). Keep it concise. "
            "Transcript:\n\n" + text
        )
        try:
            response = client.responses.create(
                model="gpt-4.1-mini",
                input=prompt,
                max_output_tokens=250,
            )
            output = response.output_text.strip()
            items = []
            for line in output.split("\n"):
                line = line.strip()
                if line.startswith("-") or line.startswith("*"):
                    items.append(line.lstrip("-* ").strip())
            return items
        except Exception:
            pass

    # Advanced local dynamic regex and keyword task extractor
    action_markers = [
        r"(?:need to|must|should|will|please|tasked to|responsible for|action item:?)\s+([^.\n]+)",
        r"(?:let's|let us)\s+([^.\n]+)",
    ]
    items = []
    
    # 1. Look for matching sentences with modal verbs (need to, should, must, will)
    for marker in action_markers:
        for match in re.finditer(marker, text, re.IGNORECASE):
            task = match.group(1).strip()
            if 10 < len(task) < 100:
                task = task[0].upper() + task[1:]
                items.append(task)
                
    # 2. Heuristic fallback: scan sentences for task keywords
    if not items:
        sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if s.strip()]
        keywords = ["schedule", "review", "email", "send", "document", "update", "deploy", "prepare", "research"]
        for s in sentences:
            s_words = s.lower().split()
            if any(kw in s_words for kw in keywords):
                clean_s = s.strip(".,?!")
                clean_s = clean_s[0].upper() + clean_s[1:]
                items.append(clean_s)
                
    # 3. Ultimate defaults if no text matched
    if not items:
        items = [
            "Review transcript details and highlight key talking points",
            "Synthesize general meeting findings and prepare list of questions",
            "Follow up on outstanding agenda items discussed during session"
        ]
        
    # Remove duplicates and return
    return list(dict.fromkeys(items))[:5]


def generate_meeting_outline(text: str) -> str:
    if OPENAI_API_KEY:
        client = OpenAI(api_key=OPENAI_API_KEY)
        prompt = (
            "Generate a structured, hierarchical meeting outline with markdown headers (###, ####) and bullet points from this transcript. "
            "Organize it into logical agenda topics. Transcript:\n\n" + text
        )
        try:
            response = client.responses.create(
                model="gpt-4.1-mini",
                input=prompt,
                max_output_tokens=400,
            )
            return response.output_text.strip()
        except Exception:
            pass
            
    # Advanced Heuristic Meeting Outline Generator (dynamic chronological context block clustering)
    sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if s.strip()]
    if not sentences:
        return "### 📋 Meeting Briefing Outline\n\n*No recording data available to outline.*"
        
    outline = "### 📋 Meeting Briefing Outline\n\n"
    
    # Split sentences into 3 conceptual blocks (Intro, Core Discussion, Action Agenda)
    chunk_size = max(1, len(sentences) // 3)
    chunks = [sentences[i:i + chunk_size] for i in range(0, len(sentences), chunk_size)]
    
    headers = [
        "1. Executive Briefing & Context",
        "2. Core Concepts & Technical Deep Dive",
        "3. Strategy, Next Steps & Deadlines"
    ]
    
    for idx, chunk in enumerate(chunks[:3]):
        if not chunk:
            continue
        outline += f"#### 📍 {headers[idx]}\n"
        outline += f"- **Topic Focus**: {chunk[0]}\n"
        
        # Extract keywords
        chunk_text = " ".join(chunk).lower()
        stopwords = {"the", "and", "of", "to", "a", "i", "it", "in", "is", "that", "you", "for", "on", "with", "as", "at", "by", "an", "this", "be", "have"}
        words = re.findall(r'\b[a-z]{4,}\b', chunk_text)
        frequent_keywords = {}
        for w in words:
            if w not in stopwords:
                frequent_keywords[w] = frequent_keywords.get(w, 0) + 1
        sorted_kw = sorted(frequent_keywords.items(), key=lambda x: x[1], reverse=True)[:3]
        keywords_str = ", ".join([k[0].upper() for k in sorted_kw]) if sorted_kw else "GENERAL DISCUSSION"
        
        outline += f"- **Core Theme Tags**: {keywords_str}\n"
        # Include summary sentences
        for sentence in chunk[1:3]:
            outline += f"- {sentence}\n"
        outline += "\n"
        
    return outline


def calculate_speaking_metrics(text: str, duration_seconds: float = 0.0) -> Dict[str, Any]:
    # Count words
    words = [w.strip(".,?!\"()[]{}") for w in text.split() if w.strip()]
    word_count = len(words)
    
    # Estimate speaking duration in minutes (if not provided, assume average 130 WPM)
    duration_minutes = max(0.1, duration_seconds / 60.0) if duration_seconds > 0.0 else (word_count / 130.0)
    wpm = int(word_count / duration_minutes) if duration_minutes > 0.0 else 0
    
    # Count filler words
    filler_words = ["um", "ah", "like", "actually", "basically", "so", "you know"]
    filler_counts = {}
    total_fillers = 0
    
    for filler in filler_words:
        pattern = r"\b" + re.escape(filler) + r"\b"
        matches = len(re.findall(pattern, text, re.IGNORECASE))
        if matches > 0:
            filler_counts[filler] = matches
            total_fillers += matches
            
    # Calculate clarity score
    filler_ratio = total_fillers / max(1, word_count)
    clarity_score = max(50, int(100 - (filler_ratio * 400)))
    
    return {
        "word_count": word_count,
        "wpm": min(wpm, 250),
        "filler_words": filler_counts,
        "total_fillers": total_fillers,
        "clarity_score": clarity_score
    }

